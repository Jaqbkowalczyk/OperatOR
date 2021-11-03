#!/usr/bin/env python
# -*- coding: utf-8 -*-

import time
import os
import re
import logging
import pyautogui
from PyPDF3 import PdfFileMerger
import webbrowser as web
import win32com.client as win32
from decimal import Decimal
from docx import Document
import tkinter as tk
from shapely.geometry import Polygon
from shapely.geometry import Point as shapelyPoint
from tkinter.filedialog import askdirectory, askopenfile
from tkinter import Menu, scrolledtext
from tkinter import messagebox
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.opc.exceptions import PackageNotFoundError
from docx.shared import Pt
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from numpy import ones, vstack
from numpy.linalg import lstsq
import csv
import math

logging.basicConfig(level=logging.DEBUG, format='%(asctime)s - %(levelname)s - %(message)s')
pyautogui.FAILSAFE = True
JEDNOSTKAREJESTROWA = ''
OBREB = ''
KERG = ''
FOLDER = ''
GMLFILE = "GD-13.6640.8414.2021_40901011.gml"
XYACCURACY = 2
HACCURACY = 2
ANGLEACCURACY = 4
AREAACCURACY = 0
DIVISIONPOINTS = 'pkty_podzial.txt'
MAINPARCEL = '38/1'

def find_kerg(filename):  # todo regex
    logging.debug(f'{filename}')
    if filename.split('.')[-1] == 'rtf':
        filename = convertrtftodocx(filename)
    text = str(get_text_from_doc(filename))
    logging.debug(f'{text}')
    kerg = re.findall(r"[\d.]+", text)
    #print(kerg)
    kerg = set(kerg)
    #print(kerg)


def set_kerg(kerg_value: str):
    global KERG
    KERG = kerg_value
    logging.debug(f'KERG został ustawiony na: {KERG}')


def set_jedn(jedn_value: str):
    global JEDNOSTKAREJESTROWA
    JEDNOSTKAREJESTROWA = jedn_value
    logging.debug(f'Jednostka została ustawiona na: {JEDNOSTKAREJESTROWA}')


def set_obr(obr_value: str):
    global OBREB
    OBREB = obr_value
    logging.debug(f'OBREB został ustawiony na: {OBREB}')


def set_gmlfile(gml_value: str):
    global GMLFILE
    GMLFILE = GML_value
    logging.debug(f'GMLFILE został ustawiony na: {GMLFILE}')


def ask_for_kerg():
    pass


def set_folder(folder_value: str):
    global FOLDER
    FOLDER = folder_value
    logging.debug(f'FOLDER został ustawiony na: {FOLDER}')


def open_folder():
    tk.Tk().withdraw()  # we don't want a full GUI, so keep the root window from appearing
    filename = askdirectory()  # show an "Open" dialog box and return the path to the selected folder
    return filename


def open_file():
    logging.debug('hej')
    tk.Tk().withdraw()  # we don't want a full GUI, so keep the root window from appearing
    filename = askopenfile('r')  # show an "Open" dialog box and return the path to the selected folder
    return filename


def convertrtftodocx(file):
    word = win32.Dispatch("Word.Application")
    wdFormatDocumentDefault = 16
    wdHeaderFooterPrimary = 1
    doc = word.Documents.Open(file)
    for pic in doc.InlineShapes:
        pic.LinkFormat.SavePictureWithDocument = True
    for hPic in doc.sections(1).headers(wdHeaderFooterPrimary).Range.InlineShapes:
        hPic.LinkFormat.SavePictureWithDocument = True
    file = file.split('.')[0]
    doc.SaveAs(str(file + ".docx"), FileFormat=wdFormatDocumentDefault)
    doc.Close()
    word.Quit()
    return str(file + ".docx")


def get_text_from_doc(filename):
    doc = Document(filename)
    fullText = []
    i = 0
    j = 0
    for para in doc.paragraphs:
        fullText.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                fullText.append(cell.text)
                j += 1
            i += 1
            j = 0
        i = 0
    return '\n'.join(fullText)


def find_info():
    for subdir, dirs, files in os.walk(FOLDER):
        for file in files:
            if file == "info_o_materiałach12.rtf":
                logging.debug('Znalazłem info o materiałach!')
                logging.debug(f'{os.path.join(subdir, file)}')
                return os.path.join(subdir, file)
            else:
                pass


def check_project_data():  # todo find a way to extract text from .rtf file
    """Function to check for variables used in all future functions"""
    # 1. Select project folder
    folder = open_folder()
    set_folder(folder)
    # 2. Find data from PODGiK (.gml file, info o materialach)
    # text = get_text_from_doc(find_info())

    #print(text)
    # 3. Search through .gml file to find JEDNOSTKAREJESTROWA and OBREB values
    # 4. Find KERG number
    set_kerg('666.2250.2021')
    pass


def write_report():
    """Function to write report file using given values"""
    s = "I love #stackoverflow# because #people# are very #helpful# #helpful#"
    hashtag = re.findall(r"#(\w+)#", s)  # znajdź wszystkie hashtagi w szablonie
    #print(set(hashtag))
    document = Document()
    for paragraph in document.paragraphs:
        if 'sea' in paragraph.text:
            #print(paragraph.text)
            paragraph.text = 'new text containing ocean'


def createparcelfilefromdoc(file, parcelsfile):
    """This function creates parcel text file from .docx table with parcels"""
    doc = Document(file)
    parcels = []
    i = 0
    j = 0
    newline = lambda x: x + '\n'
    output = open(os.getcwd() + '\\' + parcelsfile, 'w')
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if j == 1:
                    cell_parcels = cell.text.split(',')
                    #logging.debug(f'Text komórki: {cell.text}')
                    for parcel in cell_parcels:
                        parcel = parcel.replace(' ', '')
                        parcels.append(parcel)
                j += 1
            i += 1
            j = 0
        i = 0
    for parcel in set(parcels):
        output.write(newline(parcel))
    output.close()
    return output


def filldocxtemplate(templatefile, outputfile, owner=None):
    """function that copies docx template into an end of output file"""
    # select only paragraphs or table nodes
    template = Document(templatefile)
    get_text_from_doc(templatefile)
    try:
        output = Document(outputfile)
    except PackageNotFoundError:
        output = Document()
        output.save(outputfile)
    for child in template.element.body.xpath('w:p | w:tbl'):
        if isinstance(child, CT_P):
            paragraph = Paragraph(child, template)
            outpara = output.add_paragraph()
            s = paragraph.text
            hashtag = re.findall(r"#(\w+)#", s)
            logging.debug(f'Hashtagi: {hashtag}')
            for run in paragraph.runs:
                output_run = outpara.add_run(run.text)
                # Run's bold data
                for parcel in owner.parcels:
                    if parcel in output_run.text:
                        output_run.text = output_run.text.replace(parcel, '')
                        output_run.text = output_run.text.replace(' ,', '')
                        parcel_run = outpara.add_run(', ' + parcel)
                        parcel_run.bold = True
                        parcel_run.font.name = 'Times New Roman'

                output_run.bold = run.bold
                # Run's italic data
                output_run.italic = run.italic
                # Run's underline data
                output_run.underline = run.underline
                # Run's color data
                output_run.font.color.rgb = run.font.color.rgb
                # Run's font
                output_run.font.name = 'Times New Roman'
                output_run.font.size = run.font.size
                # Run's font data
                output_run.style.name = run.style.name
                # Paragraph's alignment data
            outpara.paragraph_format.line_spacing = 1.0
            outpara.paragraph_format.alignment = paragraph.paragraph_format.alignment
            outpara.paragraph_format.first_line_indent = paragraph.paragraph_format.first_line_indent
            outpara.paragraph_format.space_before = 5
            outpara.paragraph_format.space_after = 5


        elif isinstance(child, CT_Tbl):
            table = Table(child, template)
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        #logging.debug(f'Text komórki: {paragraph.text}')
                        hashtag = re.findall(r"#(\w+)#", paragraph.text)
                        logging.debug(f'Hashtagi: {hashtag}')
                        if len(hashtag) == 0:
                            pass
                        else:
                            for hash in hashtag:
                                if hash == 'imie':
                                    inline = paragraph.runs
                                    for i in range(len(inline)):
                                        if 'imie' in inline[i].text:
                                            logging.debug(f'Zamieniam #imie# na {owner.name}')
                                            text = inline[i].text.replace('#imie#', owner.name)
                                            inline[i].text = text
                                elif hash == 'nazwisko':
                                    inline = paragraph.runs
                                    for i in range(len(inline)):
                                        if 'nazwisko' in inline[i].text:
                                            logging.debug(f'Zamieniam #nazwisko# na {owner.surname}')
                                            text = inline[i].text.replace('#nazwisko#', owner.surname)
                                            inline[i].text = text
                                elif hash == 'adres':
                                    inline = paragraph.runs
                                    for i in range(len(inline)):
                                        if 'adres' in inline[i].text:
                                            logging.debug(f'Zamieniam #adres# na {owner.address}')
                                            text = inline[i].text.replace('#adres#', owner.address)
                                            text = text.replace(', ', ',\n')
                                            inline[i].text = text
                                elif hash == 'godzina':
                                    inline = paragraph.runs
                                    for i in range(len(inline)):
                                        if 'godzina' in inline[i].text:
                                            try:
                                                text = inline[i].text.replace('#godzina#', owner.hour)
                                                inline[i].text = text
                                                logging.debug(f'Wprowadziłem godzinę dla {owner.name} {owner.surname}:'
                                                              f' {owner.hour} ')
                                            except:
                                                pass

            paragraph = output.add_paragraph()
            paragraph._p.addnext(table._tbl)

            paragraph.paragraph_format.first_line_indent = 1
            paragraph.paragraph_format.space_before = 1
            paragraph.paragraph_format.space_after = 1
            paragraph.paragraph_format.line_spacing = 1
    paragraph = output.add_paragraph()
    run = paragraph.add_run()
    run.add_break(WD_BREAK.PAGE)
    output.save(outputfile)


def findowners(file, parcelspath):
    owners = []
    ownersobj = []
    parcel = ''
    doc = Document(file)
    table = doc.tables[0]
    i = 0
    j = 0
    parcelsfile = open(parcelspath, 'r')
    # find target parcels from file
    parcels = [line.replace('\n', '') for line in parcelsfile.readlines()]
    parcels = set(parcels)
    for row in table.rows:
        if i < 1:
            i += 1
            continue
        elif row.cells[0].text == 'Nr działki' or row.cells[0].text == '':
            i += 1
            continue
        # Find parcel name and chceck with target parcels
        elif row.cells[0].text.split('.')[-1] in parcels:
            for cell in row.cells:
                if j == 0:
                    parcel = cell.text.split('.')
                    parcel = parcel[-1]
                    #print(parcel)
                if j == 6:
                    #logging.debug(f'Text komórki: {cell.text}')
                    temp = cell.text.split('udział ')
                    temp.pop(0)
                    for udz in temp:
                        temp = udz.split(';')
                        temp[0] = temp[0].replace('Własność: ', '')
                        temp[1] = temp[1].replace('Własność: ', '')
                        logging.debug(f'{temp}')
                        adr = temp[1].replace('\n', '')
                        adr = adr.replace('Władanie: Użytkowanie', '')
                        if temp[0].split(' ')[1] == 'Małż.:':
                            logging.debug(f'Małżeństwo split: {temp}')
                            adr1 = temp[1].split('\n')[0]
                            adr2 = temp[2].split('\n')[0]
                            for item in temp:
                                item = item.replace('\nWłasność: ', '')
                                if '\n' in item:
                                    temp1 = item.split('\n')
                                    temp1.pop(0)
                                    temp1 = temp1[0].split(',')
                                    temp1 = temp1[0].split(' ')
                                    if len(temp1) == 3:
                                        logging.debug(f'Kasuje imię {temp1[1]}')
                                        temp1.pop(1)
                                        logging.debug(f'Współwlasciciel 2: {temp1}  adres: {adr2}')
                                        owners.append((temp1[0], temp1[1], adr2, parcel))
                                    else:
                                        logging.debug(f'Współwlasciciel 2: {temp1} adres: {adr2}')
                                        owners.append(
                                            (temp1[0], ' '.join(temp1[i] for i in range(1, len(temp1))), adr2, parcel))
                            temp = temp[0].split(',')
                            temp = temp[0].split(' ')
                            temp.pop(0)
                            temp.pop(0)
                            if len(temp) == 3 and temp[1][0].isupper():
                                logging.debug(f'Kasuje imię {temp[1]}')
                                temp.pop(1)
                                logging.debug(f'Wspolwlasciciel 1: {temp} adres: {adr1}')
                                owners.append((temp[0], temp[1], adr1, parcel))
                            else:
                                logging.debug(f'Wspolwlasciciel 1: {temp} adres: {adr1}')
                                owners.append((temp[0], ' '.join(temp[i] for i in range(1, len(temp))), adr1, parcel))

                        else:
                            temp = temp[0].split(',')
                            temp = temp[0].split(' ')
                            temp.pop(0)
                            if len(temp) == 3 and temp[1][0].isupper():
                                logging.debug(f'Kasuje imię {temp[1]}')
                                temp.pop(1)
                                logging.debug(f'Wlasciciel: {temp}')
                                owners.append((temp[0], temp[1], adr, parcel))
                            else:
                                logging.debug(f'Wlasciciel: {temp}')
                                owners.append((temp[0], ' '.join(temp[i] for i in range(1, len(temp))), adr, parcel))
                j += 1
            j = 0
        else:
            i += 1
            continue
    i = 0

    for owner in owners:
        alreadyin = False
        if len(ownersobj) == 0:
            o = Owner(owner[0], owner[1], owner[2], owner[3])
            ownersobj.append(o)
        else:
            for obj in ownersobj:
                if obj.name == owner[0] and obj.surname == owner[1] and obj.address == owner[2] and owner[3] not in obj.parcels:
                    obj.addparcels(owner[3])
                    logging.debug(f'dodano parcele do obiektu: {obj.name} {obj.surname}')
                    alreadyin = True
            if not alreadyin:
                o = Owner(owner[0], owner[1], owner[2], owner[3])
                ownersobj.append(o)

    for o in ownersobj:
        print(f'Cześć nazywam się: {o.name} {o.surname}. Mieszkam przy ul.: {o.address},'
              f' jestem właścicielem działki: {o.parcels}')
    return ownersobj


def findkw(file, parcelspath):
    "find KW for specified parcels from .docx file"
    kw = []
    doc = Document(file)
    table = doc.tables[0]
    i = 0
    j = 0
    parcelsfile = open(parcelspath, 'r')
    # find target parcels from file
    parcels = [line.replace('\n', '') for line in parcelsfile.readlines()]
    parcels = set(parcels)
    for row in table.rows:
        if i < 1:
            i += 1
            continue
        elif row.cells[0].text == 'Nr działki' or row.cells[0].text == '':
            i += 1
            continue
        # Find parcel name and chceck with target parcels
        elif row.cells[0].text.split('.')[-1] in parcels:
            for cell in row.cells:
                if j == 0:
                    parcel = cell.text.split('.')
                    parcel = parcel[-1]
                    #print(parcel)
                if j == 7:
                    #logging.debug(f'Text komórki: {cell.text}')
                    text = cell.text.replace('\t', '')
                    logging.debug(f'{text}')
                    kw.append(calccontrolnumber(text))
                j += 1
            j = 0
    #print(set(kw))
    return set(kw)


def calccontrolnumber(kw):
    """Calculate control number for specified KW"""
    dictionary = {'0': 0, '1': 1, '2': 2, '3': 3, '4': 4, '5': 5, '6': 6, '7': 7, '8': 8, '9': 9, 'X': 10, 'A': 11,
                  'B': 12, 'C': 13, 'D': 14, 'E': 15, 'F': 16, 'G': 17, 'H': 18, 'I': 19, 'J': 20, 'K': 21, 'L': 22,
                  'M': 23, 'N': 24, 'O': 25, 'P': 26, 'R': 27, 'S': 28, 'T': 29, 'U': 30, 'W': 31, 'Y': 32, 'Z': 33}
    waga = '137137137137'
    sum = 0
    i = 0
    parts = []
    if len(kw.split('/')) == 3:
        return kw
    elif len(kw.split('/')) == 2:
        parts = kw.split('/')
        if len(kw.split('/')[1]) < 8:
            parts[1] = ('0' * (8 - len(parts[1]))) + parts[1]
        kw = parts[0] + parts[1]
    else:
        kw = kw.replace(' ', '')
        kw = kw.replace('KW', '')
        if len(kw) < 8:
            kw = ('0' * (8 - len(kw))) + kw
            kw = 'KR1P' + kw
    for s in kw:
        sum += dictionary[s.upper()] * int(waga[i])
        i += 1
    i = 0
    kw = kw[:4] + '/' + kw[4:] + '/' + str(sum % 10)
    return kw


def namestofile(owners, filename):
    """Create names doc from owners data"""
    i = 0
    doc = Document()
    table = doc.add_table(rows=1, cols=4)
    table.rows[0].cells[0].text = 'Lp.'
    table.rows[0].cells[1].text = 'Imię i Nazwisko'
    table.rows[0].cells[2].text = 'Adres'
    table.rows[0].cells[3].text = 'Numer przesyłki'
    for owner in owners:
        row = table.add_row()
        row.cells[0].text = str(i)
        row.cells[1].text = owner.fullname
        row.cells[2].text = owner.address
        i += 1
    doc.save(filename)


def isthesameowner(owner1, owner2):
    if owner1.name == owner2.name and owner1.surname == owner2.surname and owner1.address == owner2.address:
        return True
    else:
        return False


def removefromlist(list, determine):
    return list


def removeduplicates(file1, file2, outputfile):
    """Remove duplicates from parcel files"""
    parcels = []
    parcelsfile = open(file1, 'r')
    parcelsw = [line.replace('\n', '') for line in parcelsfile.readlines()]
    parcelsw = set(parcelsw)
    parcelsfile.close()
    parcelsfile = open(file2, 'r')
    parcelsu = [line.replace('\n', '') for line in parcelsfile.readlines()]
    parcelsu = set(parcelsu)
    parcelsfile.close()
    for parcel in parcelsw:
        if parcel in parcelsu:
            logging.debug(f'Wywalam: {parcel}')
            continue
        else:
            parcels.append(parcel)
    output = open(outputfile, 'w')
    newline = lambda x: x + '\n'
    for parcel in parcels:
        output.write(newline(parcel))
    output.close()


def createstickers(file, outfile):
    """Create stickers doc from names doc for each owner to put on letter"""
    out = Document()
    stickerstbl = out.add_table(rows=0, cols=1)
    section = out.sections[0]
    sectPr = section._sectPr
    cols = sectPr.xpath('./w:cols')[0]
    cols.set(qn('w:num'), '3')
    doc = Document(file)
    table = doc.tables[0]
    for row in table.rows[1:]:
        text = row.cells[1].text + '\n' + row.cells[2].text
        strow = stickerstbl.add_row()
        strow.cells[0].text = text
    out.save(outfile)


def parcelfinder(parcel, ownersfile):
    with open(ownersfile, 'r', newline='') as csvfile:
        reader = csv.reader(csvfile, delimiter=',')
        for row in reader:
            if parcel in row[3]:
                print(' '.join([row[0], row[1], row[2]]))


def changehash(templatefile, outputfile, hashdict):
    template = Document(templatefile)
    get_text_from_doc(templatefile)
    try:
        output = Document(outputfile)
    except PackageNotFoundError:
        output = Document()
        output.save(outputfile)
    current_section = template.sections[-1]
    new_height, new_width = current_section.page_height, current_section.page_width
    new_section = output.sections[-1]
    new_section.orientation = current_section.orientation
    new_section.page_width = current_section.page_width
    new_section.page_height = current_section.page_height
    new_section.top_margin = current_section.top_margin
    new_section.bottom_margin = current_section.bottom_margin
    new_section.left_margin = current_section.left_margin
    new_section.right_margin = current_section.right_margin
    new_section.header.paragraphs[0].text = current_section.header.paragraphs[0].text

    for child in template.element.body.xpath('w:p | w:tbl'):
        if isinstance(child, CT_P):
            paragraph = Paragraph(child, template)
            outpara = output.add_paragraph()
            s = paragraph.text
            hashtag = re.findall(r"#(\w+)#", s)
            inline = []
            i = 0
            #logging.debug(f'Text komórki: {paragraph.text}')
            if len(hashtag) == 0:
                pass
            else:
                logging.debug(f'Hashtagi: {hashtag}')
                for hash in hashtag:
                    try:
                        inline = paragraph.runs
                        for i in range(len(inline)):
                            inline[i].text = inline[i].text.replace('#', '')
                            if hash in inline[i].text:
                                if hashdict[hash] != 'brak':
                                    text = inline[i].text.replace(hash, hashdict[hash])
                                    inline[i].text = text
                                else:
                                    inline[i].text = ''
                    except KeyError:
                        inline[i].text = ''
            for run in paragraph.runs:
                output_run = outpara.add_run(run.text)
                output_run.bold = run.bold
                # Run's italic data
                output_run.italic = run.italic
                # Run's underline data
                output_run.underline = run.underline
                # Run's color data
                output_run.font.color.rgb = run.font.color.rgb
                # Run's font
                output_run.font.name = 'Times New Roman'
                output_run.font.size = run.font.size
                # Run's font data
                output_run.style.name = run.style.name
                # Paragraph's alignment data
            outpara.paragraph_format.line_spacing = 0.8
            outpara.paragraph_format.alignment = paragraph.paragraph_format.alignment
            outpara.paragraph_format.first_line_indent = paragraph.paragraph_format.first_line_indent
            outpara.paragraph_format.space_before = 0
            outpara.paragraph_format.space_after = 0
        elif isinstance(child, CT_Tbl):
            table = Table(child, template)
            inline = []
            i = 0
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        #logging.debug(f'Text komórki: {paragraph.text}')
                        hashtag = re.findall(r"#(\w+)#", paragraph.text)
                        if len(hashtag) == 0:
                            pass
                        else:
                            logging.debug(f'Hashtagi: {hashtag}')
                            for hash in hashtag:
                                try:
                                    inline = paragraph.runs
                                    for i in range(len(inline)):
                                        inline[i].text = inline[i].text.replace('#', '')
                                        logging.debug(f'inine run: {inline[i].text}')
                                        if hash in inline[i].text:
                                            if hashdict[hash] != 'brak':
                                                text = inline[i].text.replace(hash, hashdict[hash])
                                                inline[i].text = text
                                            else:
                                                inline[i].text = ''

                                except KeyError:
                                    inline[i].text = ''
                    paragraph = output.add_paragraph()
                    paragraph.paragraph_format.first_line_indent = 0
                    paragraph.paragraph_format.space_before = 0
                    paragraph.paragraph_format.space_after = 0
                    paragraph.paragraph_format.line_spacing = 0
                    paragraph._p.addnext(table._tbl)

    output.save(outputfile)


def kwtopdf(kw):
    openweb('https://ekw.ms.gov.pl/eukw_ogol/KsiegiWieczyste')
    if pylocate(os.getcwd() + '\\img\\' + 'kw.png') is not None:
        pyautogui.click(pylocate(os.getcwd() + '\\img\\' + 'kw.png'))
        pyautogui.write(kw.split('/')[0])
        pyautogui.move(40, 0)
        pyautogui.click()
        pyautogui.write(kw.split('/')[1])
        pyautogui.move(100, 0)
        pyautogui.click()
        pyautogui.write(kw.split('/')[2])
        pyautogui.click(pylocate(os.getcwd() + '\\img\\' + 'wyszukaj.png'))
        pyautogui.click(pylocate(os.getcwd() + '\\img\\' + 'przegladanie.png'))
        time.sleep(.5)
        pyautogui.hotkey('ctrlleft', 'p')
        pyautogui.click(pylocate(os.getcwd() + '\\img\\' + 'zapisz.png'))
        time.sleep(.5)
        pyautogui.write(kw.replace('/', '_') + '_1')
        pyautogui.click(pylocate(os.getcwd() + '\\img\\' + 'zapisz2.png'))
        time.sleep(2)
        pyautogui.click(pylocate(os.getcwd() + '\\img\\' + 'dzial2.png'))
        time.sleep(.5)
        pyautogui.hotkey('ctrlleft', 'p')
        pyautogui.click(pylocate(os.getcwd() + '\\img\\' + 'zapisz.png'))
        pyautogui.write(kw.replace('/', '_') + '_2')
        pyautogui.click(pylocate(os.getcwd() + '\\img\\' + 'zapisz2.png'))
        time.sleep(.5)
        pyautogui.hotkey('ctrlleft', 'w')


def pdfmerge(folder):
    for subdir, dirs, files in os.walk(folder):
        for file in files:
            file = file.replace('.pdf', '')
            # print(file)
            if file.split('_')[-1] == '1':
                for f2 in files:
                    if f2.split('.')[-1] != 'pdf':
                        continue
                    f2 = f2.replace('.pdf', '')
                    # logging.debug(f'{f2.split("_")[1]} vs {file.split("_")[1]} and {file.split("_")[-1]}')
                    if f2.split('_')[1] == file.split('_')[1] and f2.split('_')[-1] != '1':
                        # logging.debug('hej')
                        merger = PdfFileMerger()
                        input1 = open(folder + '/' + file + '.pdf', 'rb')
                        input2 = open(folder + '/' + f2 + '.pdf', 'rb')
                        merger.append(input1, pages=(0, 1))
                        merger.append(input2)
                        name = file[:-2] + '.pdf'
                        output = open(name, "wb")
                        merger.write(output.name)



def openweb(url):
    """open desired url in basic browser"""
    web.open_new(url)


def pylocate(img):
    logging.debug(f'{img}')
    i = 0
    s = None
    while i < 5:
        try:
            logging.debug(f'próbuję znaleźć obraz')
            s = pyautogui.locateOnScreen(img, confidence=0.9)
            logging.debug(f'Znalazłem: {s}')
            if s is not None:
                return s
        except:
            time.sleep(3)
        i += 1
    return s


def getfeaturesfromgml(gmlfile, feature):
    # function that loops through gml file and return features of choice
    content = gmlfile.read()
    contentlist = content.split('<' + feature)[1:]
    for i, item in enumerate(contentlist):
        contentlist[i] = item.split('</gml:featureMember>')[0]
    """for item in contentlist:
        logging.debug(item)
        logging.debug('\n_____________________________________')"""
    return contentlist


def getcontentfromtags(text, tag):
    # Function to grab everything in between two tags, nested tags included
    try:
        content = text.split(tag)[1]  # grab text starting from tag
    except IndexError:
        content = 'brak'
        return content
    if 'xsi:nil="true"' in content.split('<')[0]:
        content = 'brak'
        return content
    content = re.sub(r'^.*?>', '', content)  # delete full tag
    content = content.split('</')[:-1]  # delete exit tag
    content = '</'.join(content)  # join every nested tags
    return content


def getinfofromtags(text, tag):
    # Funtion to get meta data from single tags (info) --> <tag/>
    tags = text.split(tag)[1:]
    info = {}
    for item in tags:
        content = item.split('/>')[0]
        content = content.split(' ')[1:]
        for meta in content:
            key = meta.split('=')[0]
            data = meta.split('=')[1]
            data = data.replace('"', '')
            info[key] = data
    return info


def populate_points_from_gml():
    # Function to search for points in gml file and create Point object for each one.
    gmlfile = open(GMLFILE, encoding='utf-8')
    pointlist = getfeaturesfromgml(gmlfile, 'egb:EGB_PunktGraniczny')
    pointsobj = []
    for point in pointlist:
        id = getcontentfromtags(point, 'egb:idPunktu')
        number = id.split('.')[-1]
        gmlid = getcontentfromtags(point, 'bt:lokalnyId')
        coordinates = getcontentfromtags(point, 'gml:pos')
        x = float(coordinates.split(' ')[0])
        y = float(coordinates.split(' ')[1])
        try:
            zrd = int(getcontentfromtags(point, 'egb:zrodloDanychZRD'))
        except ValueError:
            zrd = 'brak'
        try:
            bpp = int(getcontentfromtags(point, 'egb:bladPolozeniaWzgledemOsnowy'))
        except ValueError:
            bpp = 'brak'
        try:
            stb = int(getcontentfromtags(point, 'egb:kodStabilizacji'))
        except ValueError:
            stb = 'brak'
        try:
            rzg = int(getcontentfromtags(point, 'egb:kodRzeduGranicy'))
        except ValueError:
            rzg = 'brak'

        operat = getcontentfromtags(point, 'egb:dodatkoweInformacje')
        operat = operat.replace('operat punktu:', '')
        operat = operat.split(' ')[0]

        new_point = Point(id, number, x, y, gmlid, zrd, bpp, stb, rzg, operat)
        pointsobj.append(new_point)
        logging.debug(f'Utworzyłem nowy punkt o numerze: {new_point.number} id {new_point.point_id}, '
                      f'gmlid {new_point.gmlid}, wsp: {new_point.x}, {new_point.y}, {new_point.zrd} {new_point.bpp}'
                      f' {new_point.stb} {new_point.rzg}, operat: {new_point.operat}')
    gmlfile.close()
    return pointsobj


def populate_parcels_from_gml(pointsobj, ownersobj):
    # function to search through gml file looking for parcels, and then create Parcel object for each parcel.
    gmlfile = open(GMLFILE, encoding='utf-8')
    parcellist = getfeaturesfromgml(gmlfile, 'egb:EGB_DzialkaEwidencyjna')
    gmlfile.seek(0)
    jrglist = getfeaturesfromgml(gmlfile, 'egb:EGB_JednostkaRejestrowaGruntow')
    parcelsobj = []
    gmlfile.seek(0)
    landcatlist = getfeaturesfromgml(gmlfile, 'egb:EGB_Klasouzytek')
    gmlfile.seek(0)
    sharelist = getfeaturesfromgml(gmlfile, 'egb:EGB_UdzialWlasnosci')
    gmlfile.seek(0)
    sharelist2 = getfeaturesfromgml(gmlfile, 'egb:EGB_UdzialWeWladaniuNieruchomosciaSPLubJST')
    sharelist = sharelist + sharelist2
    gmlfile.seek(0)
    marriagelist = getfeaturesfromgml(gmlfile, 'egb:EGB_Malzenstwo')
    gmlfile.seek(0)
    institutionlist = getfeaturesfromgml(gmlfile, 'egb:EGB_Instytucja')
    for parcel_text in parcellist:
        owners = []
        id = getcontentfromtags(parcel_text, 'idDzialki')
        number = id.split('.')[-1]
        gmlid = getcontentfromtags(parcel_text, 'bt:lokalnyId')
        area = float(getcontentfromtags(parcel_text, 'egb:powierzchniaEwidencyjna'))
        # get jrg id from xlink:href tag. It's the last part of link that matters
        jrg_link = getinfofromtags(parcel_text, 'egb:JRG2')['xlink:href'].split(':')[-1]
        landcat_links = []
        for item in parcel_text.split('egb:klasouzytekWGranicachDzialki')[1:]:
            item = 'egb:klasouzytekWGranicachDzialki' + item
            item = item.split('/>')[0]
            landcat_links.append(getinfofromtags(item, 'egb:klasouzytekWGranicachDzialki')['xlink:href'].split(':')[-1])
        jrg = ''
        landcat = []
        for landcat_text in landcatlist:
             for landcat_link in landcat_links:
                if getcontentfromtags(landcat_text, 'bt:lokalnyId') == landcat_link:
                    class_id = getcontentfromtags(landcat_text, 'bt:lokalnyId')
                    ofu = getcontentfromtags(landcat_text, 'egb:OFU')
                    ozu = getcontentfromtags(landcat_text, 'egb:OZU')
                    ozk = getcontentfromtags(landcat_text, 'egb:OZK')
                    landcat_area = float(getcontentfromtags(landcat_text, 'egb:powierzchniaEwidencyjnaKlasouzytku'))
                    new_landcat = Landcat(class_id, ofu, ozu, ozk, landcat_area)
                    landcat.append(new_landcat)
                    logging.debug(f'new landcat: {new_landcat.ofu}, {new_landcat.ozu}{new_landcat.ozk} powierzchnia: {new_landcat.area}')
        for jrg_text in jrglist:
            if getcontentfromtags(jrg_text, 'bt:lokalnyId') == jrg_link:
                jrg = getcontentfromtags(jrg_text, 'egb:idJednostkiRejestrowej')
        for share_text in sharelist:
            try:
                if getinfofromtags(share_text, 'egb:JRG')['xlink:href'].split(':')[-1] == jrg_link:
                    entity = getcontentfromtags(share_text, 'egb:EGB_Podmiot')
                    counter = getcontentfromtags(share_text, 'egb:licznikUlamkaOkreslajacegoWartoscUdzialu')
                    denominator = getcontentfromtags(share_text, 'egb:mianownikUlamkaOkreslajacegoWartoscUdzialu')
                    if 'egb:osobaFizyczna5' in entity:
                        for owner in ownersobj:
                            if owner.id == getinfofromtags(entity, 'egb:osobaFizyczna5')['xlink:href'].split(':')[-1]:
                                owners.append((counter + '/' + denominator, [owner]))
                    elif 'egb:malzenstwo4' in entity:
                        marriage_link = getinfofromtags(entity, 'egb:malzenstwo4')['xlink:href'].split(':')[-1]
                        for marriage_text in marriagelist:
                            if getcontentfromtags(marriage_text, 'bt:lokalnyId') == marriage_link:
                                owner1 = ''
                                owner2 = ''
                                for owner in ownersobj:
                                    if owner.id == getinfofromtags(marriage_text, 'egb:osobaFizyczna2')['xlink:href'].split(':')[-1]:
                                        owner1 = owner
                                    if owner.id == getinfofromtags(marriage_text, 'egb:osobaFizyczna3')['xlink:href'].split(':')[-1]:
                                        owner2 = owner
                                if owner1 and owner2 == '':
                                    owners = None
                                else:
                                    owners.append((counter + '/' + denominator, [owner1, owner2]))
                    elif 'egb:instytucja3' in entity:
                        institution_link = getinfofromtags(entity, 'egb:instytucja3')['xlink:href'].split(':')[-1]
                        for institution_text in institutionlist:
                            if getcontentfromtags(institution_text, 'bt:lokalnyId') == institution_link:
                                for owner in ownersobj:
                                    if owner.id == institution_link:
                                        owners.append((counter + '/' + denominator, [owner]))
            except KeyError:
                pass
        kw = getcontentfromtags(parcel_text, 'egb:numerElektronicznejKW')
        if kw == 'brak':
            kw = getcontentfromtags(parcel_text, 'egb:numerKW')
        if kw != 'brak':
            kw = calccontrolnumber(kw)

        poslist = getcontentfromtags(parcel_text, 'gml:posList')
        poslist = poslist.split(' ')
        if poslist[0] == '':
            poslist.pop(0)
        pointslist = []
        points = []
        for i, coordinate in enumerate(poslist):
            if len(poslist[i]) == 0:
                continue
            if i % 2 == 0 and i<len(poslist):
                try:
                    pointslist.append((float(poslist[i]), float(poslist[i + 1])))  # Only add correct points
                    logging.debug(f'Punkt: {(float(poslist[i]), float(poslist[i + 1]))}')
                except ValueError:
                    print(number + ' ' + poslist[i] + ' ' + poslist[i+1])
                except IndexError:
                    print(number + ' ' + str(i) + str(len(poslist)))
        pointslist.pop()  # remove last point because it's the same as first
        for point in pointslist:
            for pt in pointsobj:
                if pt.x == point[0] and pt.y == point[1]:  # find point objects basing on coordinates
                    points.append(pt)

        new_parcel = Parcel(id, number, points, landcat, gmlid, area, jrg, owners=owners, kw=kw)
        new_parcel.calc_area = new_parcel.calculate_area()
        parcelsobj.append(new_parcel)
        logging.debug(f'Utworzyłem nową działkę o numerze: {new_parcel.number} id {new_parcel.parcel_id}, '
                      f'gmlid {new_parcel.gmlid}, powierzchni {new_parcel.area}, '
                      f'powierzchnia obliczona: {new_parcel.calc_area}, '
                      f'Jednostka Rejestracji Gruntów: {new_parcel.jrg}, nr KW: {new_parcel.kw} '
                      f'Właściciel: {new_parcel.get_owners()}')
    gmlfile.close()
    return parcelsobj


def populate_owners_from_gml():
    # function to search through gml file looking for parcels, and then create Parcel object for each parcel.
    gmlfile = open(GMLFILE, encoding='utf-8')
    ownerslist = getfeaturesfromgml(gmlfile, 'egb:EGB_OsobaFizyczna')
    # print(ownerslist)
    gmlfile.seek(0)
    addresslist = getfeaturesfromgml(gmlfile, 'egb:EGB_Adres')
    ownersobj = []
    address = ''
    iaddress = ''
    for owner_text in ownerslist:
        # print(owner_text)
        id = getcontentfromtags(owner_text, 'bt:lokalnyId')
        name = getcontentfromtags(owner_text, 'egb:pierwszeImie')
        name2 = getcontentfromtags(owner_text, 'egb:drugieImie')
        if name2 == 'brak':
            name2 = None
        surname = getcontentfromtags(owner_text, 'egb:pierwszyCzlonNazwiska')

        if surname == 'brak':
            surname = None
        surname2 = getcontentfromtags(owner_text, 'egb:drugiCzlonNazwiska')
        if surname2 == 'brak':
            surname2 = None
        pesel = getcontentfromtags(owner_text, 'egb:pesel')
        fathername = getcontentfromtags(owner_text, 'egb:imieOjca')
        mothername = getcontentfromtags(owner_text, 'egb:imieMatki')

        # get address id from xlink:href tag. It's the last part of link that matters
        try:
            address_link = getinfofromtags(owner_text, 'egb:adresOsobyFizycznej')['xlink:href'].split(':')[-1]
        except KeyError:
            address = 'brak'
            address_link = ''
        for address_text in addresslist:
            if getcontentfromtags(address_text, 'bt:lokalnyId') == address_link:
                country = getcontentfromtags(address_text, 'egb:kraj')
                code = getcontentfromtags(address_text, 'egb:kodPocztowy')
                town = getcontentfromtags(address_text, 'egb:miejscowosc')
                number = getcontentfromtags(address_text, 'egb:numerPorzadkowy')
                localnum = getcontentfromtags(address_text, 'egb:nrLokalu')
                street = getcontentfromtags(address_text, 'egb:ulica')
                if country =='brak':
                    country = ''
                if localnum != 'brak':
                    address = street + ' ' + number + '/' + localnum + '\n' + code + ' ' + town + ' ' + country
                elif street != 'brak':
                    address = street + ' ' + number + '\n' + code + ' ' + town + ' ' + country
                else:
                    address = town + ' ' + number + '\n' + code + ' ' + town + ' ' + country
        new_owner = Owner(id, name, address, surname=surname, pesel=pesel, fathername=fathername, mothername=mothername,
                          name2=name2, surname2=surname2)
        ownersobj.append(new_owner)
        logging.debug(f'Utworzyłem nowego właściciela: {new_owner.name} {new_owner.surname} id {new_owner.id},\n'
                      f'PESEL {new_owner.pesel}, imiona rodziców: {new_owner.fathername} i {new_owner.mothername},\n'
                      f'Adres: {new_owner.address}')
    gmlfile.seek(0)
    institutionlist = getfeaturesfromgml(gmlfile, 'egb:EGB_Instytucja')
    for institution_text in institutionlist:
        id = getcontentfromtags(institution_text, 'bt:lokalnyId')
        name = getcontentfromtags(institution_text, 'nazwaPelna')
        regon = getcontentfromtags(institution_text, 'egb:regon')
        nip = getcontentfromtags(institution_text, 'egb:nip')
        try:
            address_link = getinfofromtags(institution_text, 'egb:adresInstytucji')['xlink:href'].split(':')[-1]
        except KeyError:
            iaddress = 'brak'
            address_link = ''
        for address_text in addresslist:
            if getcontentfromtags(address_text, 'bt:lokalnyId') == address_link:
                country = getcontentfromtags(address_text, 'egb:kraj')
                code = getcontentfromtags(address_text, 'egb:kodPocztowy')
                town = getcontentfromtags(address_text, 'egb:miejscowosc')
                number = getcontentfromtags(address_text, 'egb:numerPorzadkowy')
                localnum = getcontentfromtags(address_text, 'egb:nrLokalu')
                street = getcontentfromtags(address_text, 'egb:ulica')
                if country =='brak':
                    country = ''
                if localnum != 'brak':
                    iaddress = street + ' ' + number + '/' + localnum + '\n' + code + ' ' + town + ' ' + country
                elif street != 'brak':
                    iaddress = street + ' ' + number + '\n' + code + ' ' + town + ' ' + country
                else:
                    iaddress = town + ' ' + number + '\n' + code + ' ' + town + ' ' + country
        new_owner = Owner(id, name, iaddress, nip=nip, regon=regon)
        ownersobj.append(new_owner)
        logging.debug(f'Utworzyłem nową instytucję: {new_owner.name} id {new_owner.id}\n'
                      f'NIP: {new_owner.nip} REGON: {new_owner.regon}\n'
                      f'Adres: {new_owner.address}')
    gmlfile.close()
    return ownersobj



def populate_points_from_csv(file):
    """Fuction used to import points from .csv file"""
    pointsobj = []
    with open(str(file)) as csvfile:
        data = {num: (float(x), float(y))
                for num, x, y in csv.reader(csvfile, delimiter=' ')}
    # print(f'Zaimportowano punkty: {data}')
    for number, coordinates in data.items():
        x, y = coordinates
        new_point = Point(number, number, x, y)
        pointsobj.append(new_point)
        logging.debug(f'Utworzyłem nowy punkt o numerze: {new_point.number} id {new_point.point_id}')
    return pointsobj


def distance_from_line(point, line):
    """Fuction provides distance from point to given line"""
    x, y = point
    # ax = by + c
    # ax - by - c = 0
    a, b, c = line
    y0 = 0
    x0 = (y0 * b + c) / a
    distance = abs(((-b) * y + a * x - c)) / (math.sqrt((a ** 2) + (b ** 2)))
    yt = y0 + 500
    xt = (yt * b + c) / a
    position = (xt - x0) * (y - y0) - (yt - y0) * (x - x0)
    if position > 0:
        sign = 1
    elif position < 0:
        sign = -1
    else:
        sign = 0
    distance = distance * -sign
    # print(f'Distance from line is: {round(distance, 3)}')
    return round(distance, 3)


def line_from_two_points(point1=(0.0, 0.0), point2=(1.0, 1.0)):
    """Fuction docstring"""
    points = [point1, point2]
    x_coords, y_coords = zip(*points)
    a = vstack([y_coords, ones(len(y_coords))]).T
    b, c = lstsq(a, x_coords, rcond=None)[0]
    # print("Line Solution is x = {b}y + {c}".format(b=b, c=c))
    return (1, b, c)


def is_on_border(parcel, point):
    parcel: Parcel
    text = 'Parcel points: '
    logging.debug(f'\n\nParcel: {parcel.number}')
    for pt in parcel.points:
        text += f'{pt.number}, '
    logging.debug(text+'\n\n')
    i = 0
    if len(parcel.points) == 0:
        return False
    while i < len(parcel.points)-1:
        border = line_from_two_points((parcel.points[i].x, parcel.points[i].y),
                                      (parcel.points[i+1].x, parcel.points[i+1].y))
        #logging.debug(f'Line from points: {parcel.points[i].number}, {parcel.points[i+1].number}')
        distance_line = abs(distance_from_line((point.x, point.y), border))
        logging.debug(f'point number: {point.number}, line: {parcel.points[i].number}, {parcel.points[i+1].number} \n'
                      f'distance from line: {distance_line}')
        if abs(distance_line) < 0.05:
            border_length = abs(math.dist([parcel.points[i].x, parcel.points[i].y],
                                      [parcel.points[i+1].x, parcel.points[i+1].y]))
            dist1 = abs(math.dist([parcel.points[i].x, parcel.points[i].y], [point.x, point.y]))
            dist2 = abs(math.dist([parcel.points[i+1].x, parcel.points[i+1].y], [point.x, point.y]))
            if border_length < dist1 or border_length < dist2:
                logging.debug(f'Point number: {point.number} is on BORDER EXTENSION not actual border')
                return False
            logging.debug(f'Point number: {point.number} is on border of parcel: {parcel.number}, on border between '
                          f'points: {parcel.points[i].number}, and {parcel.points[i+1].number}')
            return True
        else:
            pass
        i += 1
        if i == len(parcel.points)-1:
            border = line_from_two_points((parcel.points[i].x, parcel.points[i].y),
                                          (parcel.points[0].x, parcel.points[0].y))
            logging.debug(f'Line from points: {parcel.points[i].number}, {parcel.points[0].number}')
            distance = distance_from_line((point.x, point.y), border)
            logging.debug(f'point number: {point.number}, x: {point.x},y: {point.y}\n'
                          f'distance from line: {distance}')
            if abs(distance) < 0.05:
                border_length = abs(math.dist([parcel.points[i].x, parcel.points[i].y],
                                              [parcel.points[0].x, parcel.points[0].y]))
                dist1 = abs(math.dist([parcel.points[i].x, parcel.points[i].y], [point.x, point.y]))
                dist2 = abs(math.dist([parcel.points[0].x, parcel.points[0].y], [point.x, point.y]))
                if border_length < dist1 or border_length < dist2:
                    logging.debug(f'Point number: {point.number} is on BORDER EXTENSION not actual border')
                    return False
                logging.debug(
                    f'Point number: {point.number} is on border of parcel: {parcel.number}, on border between '
                    f'points: {parcel.points[i].number}, and {parcel.points[0].number}')
                return True

    return False


def list_from_csv(csvfile, delimiter=','):
    with open(csvfile, 'r', newline='') as file:
        reader = csv.reader(file, delimiter=delimiter)
        data = list(reader)
    logging.debug(f'Data: {data[0]}')
    return data[0]


def write_area_to_file(parcels, file):
    with open(file, 'w', newline='') as csvfile:
        writer = csv.writer(csvfile, delimiter=',')
        for parcel in parcels:
            writer.writerow([parcel.number, parcel.area, float(parcel.calc_area)/10000])


def write_parcel_points_to_file(parcels, file):
    with open(file, 'w', newline='') as csvfile:
        writer = csv.writer(csvfile, delimiter=',')
        for parcel in parcels:
            for point in parcel.points:
                writer.writerow([point.number, point.x, point.y])

def fill_changes_report(parcel):
    number = parcel.number.replace('/', '_')
    filename = 'Wykaz zmian ' + number + '.rtf'
    origfile = 'Wykaz zmian.docx'
    wykazdict = {'enter': '\n' * 30, "nr_dz": str(parcel.number).split('/')[0], 'kw': parcel.kw, 'owner': parcel.get_owners(), 'parcel_id': parcel.parcel_id,
                 "pow_ewid": str(parcel.area), "jed_ewid": parcel.jed_ewid, "obr": parcel.obr, "jrg": parcel.jrg}
    ha = int(parcel.area)  # round down parcel area to get ha
    a = int((parcel.area - ha) * 100)
    m2 = int((((parcel.area - ha) * 100) - a) * 100)
    ha = str(ha)
    a = str(a)
    m2 = str(m2)
    if len(a) < 2:
        a = '0' + a
    if len(m2) < 2:
        m2 = '0' + m2
        if m2 == '00':
            m2 = '--'
    wykazdict['ha'] = ha
    wykazdict['a'] = a
    wykazdict['m2'] = m2

    for i, landcat in enumerate(parcel.landcat):
        ofukey = 'ofu' + str(i)
        ozukey = 'ozu' + str(i)
        ozkkey = 'ozk' + str(i)
        areakey = 'landcat_area' + str(i)
        wykazdict[ofukey] = landcat.ofu
        wykazdict[ozukey] = landcat.ozu
        wykazdict[ozkkey] = landcat.ozk
        wykazdict[areakey] = str(landcat.area)

    changehash(os.getcwd() + '\\docs\\' + origfile, filename, hashdict=wykazdict)


class Owner:
    def __init__(self, id, name, address, surname=None, pesel=None, fathername=None, mothername=None, parcels=[],
                 name2=None, surname2=None, nip=None, regon=None, hour=None,
                 date=None, source=None):
        self.id = id
        self.name = name
        self.surname = surname
        self.address = address
        self.name2 = name2
        self.surname2 = surname2
        self.pesel = pesel
        self.fathername = fathername
        self.mothername = mothername
        self.nip = nip
        self.regon = regon
        self.hour = hour
        self.date = date
        self.source = source
        self.parcels = parcels
        try:
            self.fullname = self.name + ' ' + self.surname
        except TypeError:
            self.fullname = self.name
    def addparcels(self, parcel):
        self.parcels.append(parcel)

    def zawiadomienie(self):
        pass


class Parcel:
    def __init__(self, parcel_id, number, points, landcat, gmlid=None, area=None, jrg=None, owners=None, kw=None,
                 calc_area=None, jed_ewid=None, obr=None):
        self.parcel_id = parcel_id
        self.gmlid = gmlid
        self.number = number
        self.owners = owners
        self.points = points  # List of Point objects
        self.landcat = landcat  # List of Landcat objects
        self.kw = kw
        self.area = area
        self.jrg = jrg
        self.calc_area = calc_area
        if jed_ewid is not None:
            self.jed_ewid = jed_ewid
        else:
            self.jed_ewid = self.parcel_id.split('.')[0]
        if obr is not None:
            self.obr = obr
        else:
            self.obr = self.parcel_id.split('.')[1]
        #self.Polygon = Polygon(pointlist) <- inaczej jakoś

    def calculate_area(self):
        pointlist = []
        for pointobj in self.points:
            pointlist.append((pointobj.x, pointobj.y))
        pgon = Polygon(pointlist)
        #logging.debug(f'Parcel calculated Area: {round(pgon.area)}')
        return round(pgon.area)

    def get_owners(self):
        text = ''
        for owner in self.owners:
            if isinstance(owner[1][0], Owner):
                if owner[1][0].surname is not None:
                    if len(owner[1]) == 1:
                        text += owner[1][0].name + ' ' + owner[1][0].surname + '\nim. rodziców: ' + owner[1][0].fathername \
                                + ' i ' + owner[1][0].mothername
                    elif len(owner[1]) == 2:
                        text += 'Małż:' + ' ' + owner[1][0].name + ' ' + owner[1][0].surname + \
                               ' im. rodziców: ' + owner[1][0].fathername + ' i ' + owner[1][0].mothername + '\n' + \
                               owner[1][1].name + ' ' + owner[1][1].surname + ' im. rodziców: ' + owner[1][1].fathername + \
                                ' i ' + owner[1][1].mothername
                    else:
                        text = 'brak'
                else:
                    text += owner[1][0].name + ' Udziały w części:' + ' ' + owner[0]
                text += '\n'
        return text


class Point:
    def __init__(self, point_id, number, x, y, gmlid=None, zrd=None, bpp=None, stb=None, rzg=None, operat=None, sporna=None):
        self.point_id = point_id
        self.gmlid = gmlid
        self.number = number
        self.x = x
        self.y = y
        self.zrd = zrd
        self.bpp = bpp
        self.stb = stb
        self.rzg = rzg
        self.operat = operat
        self.sporna = sporna
        self.shapelyPoint = shapelyPoint(self.x, self.y)


class Landcat:
    def __init__(self, class_id, ofu, ozu, ozk, area):
        self.class_id = class_id
        self.ofu = ofu
        self.ozu = ozu
        self.ozk = ozk
        self.area = area
        self.classification = ozu + ozk


class Navbar(tk.Frame): ...


class Toolbar(tk.Frame): ...


class Statusbar(tk.Frame): ...


class Main(tk.Frame): ...


class MainApplication(tk.Frame):
    def __init__(self, parent, *args, **kwargs):
        tk.Frame.__init__(self, parent, *args, **kwargs)
        self.parent = parent
        self.statusbar = Statusbar(self, ...)
        self.toolbar = Toolbar(self, ...)
        self.navbar = Navbar(self, ...)
        self.main = Main(self, ...)

        self.statusbar.pack(side="bottom", fill="x")
        self.toolbar.pack(side="top", fill="x")
        self.navbar.pack(side="left", fill="y")
        self.main.pack(side="right", fill="both", expand=True)


def main():
    set_kerg('6640.9602.2021')
    set_jedn('120617_2 Zielonki')
    set_obr('0003 Bosutów')
    pointsobj = populate_points_from_gml()
    ownersobj = populate_owners_from_gml()
    parcelsobj = populate_parcels_from_gml(pointsobj, ownersobj)
    #dividepointsobj = populate_points_from_csv(DIVISIONPOINTS)
    parcels_to_divide = list_from_csv('dzialki do podzialu.txt')
    divideparcelsobj = []
    divideownersobj = []
    connection_list = []
    main_points = []
    for parcel in parcels_to_divide:
        for object in parcelsobj:
            if parcel == object.number:
                divideparcelsobj.append(object)
                if parcel == MAINPARCEL:
                    main_points = object.points
                    divideparcelsobj.pop()
                logging.debug(f'Dorzucam do działek dzielonych działkę: {parcel}')
    for parcel in divideparcelsobj:
        print(parcel.owners)
        for share in parcel.owners:
            for owner in share[1]:
                ispresent = False
                for obj in divideownersobj:
                    if obj.id == owner.id:
                        ispresent = True
                if not ispresent:
                    logging.debug(f'Dzialka: {parcel.number} Wlasciciel: {owner.fullname} Adres: {owner.address}')
                    divideownersobj.append(owner)
    #namestofile(divideownersobj, 'nazwiska i adresy.docx')
    #createstickers('nazwiska i adresy.docx', 'naklejki.docx')
    #write_area_to_file(divideparcelsobj, 'powierzchnie_ewid.csv')
    for parcel in divideparcelsobj:
        fill_changes_report(parcel)

    #write_parcel_points_to_file(temp, 'punkty_32_4.csv')
    # todo search only in divided parcels, and optimize search.
    with open('wykaz wspolrzednych.csv', 'w', newline='') as csvfile:
        writer = csv.writer(csvfile, delimiter=',')
        pointlist = []
        for parcel in divideparcelsobj:
            for point in parcel.points:
                    if point.number not in pointlist:
                        """writer.writerow([point.number, str(round(point.x, 2)), str(round(point.y, 2)),
                                         point.zrd, point.bpp,
                                         point.stb, point.rzg])"""
                        writer.writerow([point.number, str(round(point.x, 2)), str(round(point.y, 2))])
                        pointlist.append(point.number)
    # pdfmerge(open_folder())
    # Write kw to file with parcel
    """with open('dzialka_kw.csv', 'w', newline='') as csvfile:
        writer = csv.writer(csvfile, delimiter=',')
        for parcel in divideparcelsobj:
            writer.writerow([parcel.number,parcel.kw])"""
    """with open('atrybuty.csv', 'w', newline='') as csvfile:
        writer = csv.writer(csvfile, delimiter=',')
        pointlist = []
        for parcel in divideparcelsobj:
            for i, point in enumerate(parcel.points):
                if 0 < i < len(parcel.points)-1:
                    if point in main_points:
                        logging.debug(f'{parcel.number}')
                        for k in range(-1, 2):
                            if parcel.points[i+k].number not in pointlist:
                                writer.writerow([parcel.points[i+k].number,
                                                 parcel.points[i+k].zrd, parcel.points[i+k].bpp,
                                                 parcel.points[i+k].stb, parcel.points[i+k].rzg])
                                pointlist.append(parcel.points[i+k].number)
                elif i == 0:
                    if point in main_points:
                        logging.debug(f'{parcel.number}')
                        for k in [len(parcel.points)-1, 0, 1]:
                            if parcel.points[i+k].number not in pointlist:
                                writer.writerow([parcel.points[i+k].number,
                                                 parcel.points[i+k].zrd, parcel.points[i+k].bpp,
                                                 parcel.points[i+k].stb, parcel.points[i+k].rzg])
                                pointlist.append(parcel.points[i+k].number)
                elif i == len(parcel.points)-1:
                    if point in main_points:
                        logging.debug(f'{parcel.number}')
                        for k in [-1, 0, -(len(parcel.points)-1)]:
                            if parcel.points[i+k].number not in pointlist:
                                writer.writerow([parcel.points[i+k].number,
                                                 parcel.points[i+k].zrd, parcel.points[i+k].bpp,
                                                 parcel.points[i+k].stb, parcel.points[i+k].rzg])
                                pointlist.append(parcel.points[i+k].number)
    
    ^
    |
    |
    Moduł do wykazu punktów granicznych dla granic odchodzących od drogi prostopadle
"""
    """for parcel in divideparcelsobj:
        for point in dividepointsobj:
            if is_on_border(parcel, point):
                is_present = False
                for item in connection_list:
                    if item[0] == parcel:
                        item.append(point)
                        is_present = True
                        break
                if not is_present:
                    connection = [parcel, point]
                    connection_list.append(connection)
    text = ''
    for c in connection_list:
        for i in c:
            text += f'{i.number}, '
        text += '\n'
    logging.debug(text)"""


    # getinfofromtags(parcellist[0], 'gml:Point')
    """point1 = Point(1,1,0,0)
    point2 = Point(2, 2, 13, 0)
    point3 = Point(3, 3, 10, 12)
    point4 = Point(4, 4, 110, 10)
    points = [point1,point2,point3,point4]
    parc = Parcel(1,[],points,101)
    parc.calculate_area()"""
    """root = tk.Tk()
    MainApplication(root).pack(side="top", fill="both", expand=True)
    root.mainloop()"""
    # toplevel = tk.Toplevel(root)

    # create a toplevel menu
    """  menubar = tk.Menu(toplevel)
    menubar.add_command(label="Podział")
    menubar.add_command(label="Quit!", command=root.quit)
    # display the menu
    toplevel.config(menu=menubar)
    root.mainloop()"""
    """ Main program """
    """x = 400
    y = 200
    pyautogui.moveTo(x, y)"""
    # check_project_data()
    # write_report()
    # ConvertRtfToDocx('C:\\Users\\Jurek\\Documents\\Kuba\\Python\\OperatOR\\docs','info_o_materiałach1.rtf')
    # createparcelfilefromdoc(os.getcwd() + '\\docs\\protokol_wyznaczenia_granic.docx', 'parcels.txt')
    """o = Owner('Jakub', 'Kowwalczyk', 'KRakowska 23, 31-102 KRaków', '512')
        filldocxtemplate(os.getcwd() + '\\docs\\Zawiad o wyznaczeniu granic.docx', 'wyzn_granic.docx', o)"""
    # removeduplicates('parcelsw.txt', 'parcelsu.txt', 'parcelsw.txt')
    # ownersfile = open_file()
    # ownersu = findowners(ownersfile, 'parcelsu.txt')
    # ownersw = findowners(ownersfile, 'parcelsw.txt')
    # owners = ownersu + ownersw
    """with open('owners.csv', 'w', newline='') as csvfile:
        writer = csv.writer(csvfile, delimiter=',')
        for owner in ownersu:
            writer.writerow([owner.name, owner.surname, owner.address, owner.parcels])"""
    # parcelfinder('581/4', ownersfile='owners.csv')
    # namestofile(owners, 'nazwiska i adresy.docx')
    # createstickers('nazwiska i adresy.docx', 'naklejki.docx')
    # file = open_file()
    # find_kerg(file.name)
    # kwlist = findkw('C:\\Users\\Jurek\\Dysk Google\\GEO\\Bibice_Zbożowa\\PODGiK\\właściciele.docx', 'parcelsu.txt')
    # for kw in kwlist:
    #    kwtopdf(kw)
    # kwtopdf('KR1P/00516204/5')
    # pdfmerge(open_folder()) #merge first page of _1 file and _2 file

    """for owner in ownersu:
        filldocxtemplate(os.getcwd() + '\\docs\\Zawiadomienie ustalenie.docx', 'ust_granicKwiatowa.docx', owner)
    for owner in ownersw:
        filldocxtemplate(os.getcwd() + '\\docs\\Zawiadomienie ustalenie.docx', 'ust_granicKwiatowa.docx', owner)"""
    return 0


if __name__ == "__main__":
    main()
